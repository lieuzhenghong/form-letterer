import csv
import re
from docx import Document

file_name = input('Enter the file name here (must be in same directory: ')

# Read the file and put it into an array
with open('data.csv', newline='') as csvfile:
    reader = csv.reader(csvfile)
    data = []
    for idx, row in enumerate(reader):
        if idx == 0:
            headers = row
        else:
            data.append(row)

def replace_text(old, rep):
    '''
    Replaces a string with another.
    Supports cross-run replacement.
    '''
# Get the start and end of the string
    idx = para.text.find(old) 
    #Edge case here: if there are two same replacement strings, how? 
    end = idx + len(old)
    i = 0
    runs = []
    added = False
    for run in para.runs:
        prev = i;
        i += len(run.text)
        #print(run.text, i, idx, end)
        # print(run.text, 'i :' , i)
        if (i>= end):
            if (added == False):
                added = True
                # this is the first run. it starts before idx and ends after end.
                # Therefore, we must keep the before idx part and the after end part
                run.text = run.text[0:idx-prev] + rep + run.text[idx-prev+len(old):]
                break;
            elif (added == True):
                # This run contains the tail of the search string. We keep everything after end.
                run.text = run.text[end-prev:]
                break;
        if (i > idx):
            if (added == False):
                added = True
                # this is the first run. it starts before idx and ends after idx.
                run.text = run.text[0:idx-prev] + rep
                #print(run.text)
            else:
                runs.append(run)
    for run in runs:
# i have all the runs from the one just before to the one just after
        run.text = '' 
        pass

for idx, row in enumerate(data):
    document = Document('{}.docx'.format(file_name))
    for para in document.paragraphs:
        m = re.findall(r'`(.+?)`', para.text)
        for match in m:
            # Find the start and end of the substring
            rep = str(row[headers.index(match)]);
            n = '`'+match+'`'
            replace_text(n, rep)

    for para in document.paragraphs:
        m = re.findall(r'{{(.+?)}}', para.text)
        for match in m:
            #print('match:' ,match)
            #print(para.text)
            match2 = match.replace('‘',"'")
            match2 = match2.replace('’',"'")
            x = str(eval(match2))
            replace_text('{{' + match + '}}', x)
    document.save('./build/letter_{}.docx'.format(row[headers.index('name')]))
