import csv
from docx import Document

total = 3000000

# Read the file and put it into an array
with open('data.csv', newline='') as csvfile:
    document = Document('./letter.docx')
    reader = csv.reader(csvfile)
    data = []
    for idx, row in enumerate(reader):
        if idx == 0:
            headers = row
        else:
            data.append(row)
    #print(headers)
    #print(data)

    for idx, row in enumerate(data):
        document = Document('letter.docx')
        #print(idx, row)
        for paragraph in document.paragraphs:
            text = paragraph.text
            if 'title' in paragraph.text:
                for run in paragraph.runs:
                    if run.text == 'title':
                        run.text = str(row[headers.index('title')]);
            if 'name1' in paragraph.text:
                text = paragraph.text
                for run in paragraph.runs:
                    if run.text == "name1" or "name1" in run.text:
                        run.text = str(row[headers.index('name')])
            if 'address1' in paragraph.text:
                for run in paragraph.runs:
                    if run.text == 'address1' or 'address1' in run.text:
                        run.text = run.text.replace(
                            "address1",
                            str(row[headers.index('address1')]).strip()
                            )
            if 'address2' in paragraph.text:
                for run in paragraph.runs:
                    if run.text == 'address2' or 'address2' in run.text:
                        run.text = run.text.replace(
                            "address2",
                            str(row[headers.index('address2')]).strip()
                            )
            if 'address3' in paragraph.text:
                text = paragraph.text
                text = text.replace(
                    "address3",
                    str(row[headers.index('address3')]).strip()
                    )
                paragraph.text = text
            if 'amount' in paragraph.text:
                amount = '{0:,}'.format(int(row[headers.index('amount')]))
                for run in paragraph.runs:
                    if "amt" in run.text:
                        run.text = run.text.replace(
                            "amt",
                            'US${}'.format(amount) 
                        )
                    if 'percentowned' in run.text:
                        run.text = run.text.replace(
                            "percentowned",
                            '{0:,}'.format(round(int(row[headers.index('amount')])/total * 100, 2))
                        )
                    if 'amountshares' in run.text:
                        run.text = run.text.replace(
                            "amountshares",
                            str(int(500000 * int(row[headers.index('amount')])/total))
                        )
            # All checks complete
        document.save('./build/letter_{}.docx'.format(row[headers.index('name')]))

