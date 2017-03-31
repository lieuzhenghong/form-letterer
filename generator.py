import csv
import re
from docx import Document


def replace_text(para, old, rep):
    '''
    Replaces a string with another.
    Supports cross-run replacement.
    '''
# Get the start and end of the string
    idx = para.text.find(old)
    end = idx + len(old)
    i = 0
    runs = []
    added = False
    for run in para.runs:
        prev = i
        i += len(run.text)
        if (i >= end):
            if added is False:
                added = True
                # this is the first run. it starts before idx and ends after
                # end.  Therefore, we must keep the before idx part and the
                # after end part
                run.text = run.text[0:idx-prev] + \
                    rep + \
                    run.text[idx-prev+len(old):]
                break
            elif added is True:
                # This run contains the tail of the search string. We keep
                # everything after end.
                run.text = run.text[end-prev:]
                break
        if (i > idx):
            if added is False:
                added = True
                # this is the first run. it starts before idx and ends after
                # idx. we should replace this first run with rep and remove
                # the rest
                run.text = run.text[0:idx-prev] + rep
            else:
                # these are the runs that are not needed; they lie completely
                # within idx and end
                run.text = ''


def generate_letters(csv_filename, letter_filename):
    try:
        # Read the file and put it into an array
        try:
            with open('{}.csv'.format(csv_filename), newline='') as csvfile:
                reader = csv.reader(csvfile)
                data = []
                for idx, row in enumerate(reader):
                    if idx == 0:
                        headers = row
                    else:
                        data.append(row)
        except PermissionError as error:
            return('<b> PERMISSION ERROR: ' + str(error) + '</b><p>You have a'
                   'csv file open and so the program cannot run. Make sure to'
                   'close all' 'csv files before running the program.')
        except FileNotFoundError as error:
            return('<b>FILE NOT FOUND ERROR</b>:'
                   'File {}.csv not found.'.format(csv_filename))
        for idx, row in enumerate(data):
            try:
                document = Document('{}.docx'.format(letter_filename))
            except FileNotFoundError as error:
                return('ERROR: File {}.docx not found.'.format(letter_filename))
            except PermissionError as error:
                return('<b>PERMISSION ERROR: ' + str(error) + '</b><p>'
                       'You have a'
                       'file open and so the program cannot run. Make sure to'
                       'close\ all docx files before running the program.')
            for para in document.paragraphs:
                m = re.findall(r'`(.+?)`', para.text)
                for match in m:
                    # Find the start and end of the substring
                    try:
                        rep = str(row[headers.index(match)])
                    except ValueError as error:
                        return('<b>VALUE ERROR: ' + str(error) + '</b><p>This'
                               'often\ means that your word document has a'
                               'thing surrounded\ by backticks that does not'
                               'exist in your CSV. Check\ again.')
                    n = '`'+match+'`'
                    replace_text(para, n, rep)

            for para in document.paragraphs:
                m = re.findall(r'{{(.+?)}}', para.text)
                for match in m:
                    match2 = match.replace('‘', "'")
                    match2 = match2.replace('’', "'")
                    x = str(eval(match2))
                    replace_text(para, '{{' + match + '}}', x)
            try:
                document.save('./build/{0}_{1}.docx'.format(letter_filename,
                              row[headers.index('name')]))
            except PermissionError as error:
                return('<b>PERMISSION ERROR: ' + str(error) + '</b><p>You have'
                       'a\ file open and so the program cannot run. Make sure'
                       'to close\ all docx files before running the program.')
        return('OK')
    except Exception as e:
        return(str(e))

if __name__ == '__main__':
    import sys
    generate_letters(sys.argv[1], sys.argv[2])
