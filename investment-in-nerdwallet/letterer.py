import csv
from docx import Document

with open('data.csv', newline='') as csvfile:
    document = Document('./form_letter.docx')
    print(document.styles['Normal'])
    reader = csv.reader(csvfile)
    data = []
    for idx, row in enumerate(reader):
        if idx == 0:
            headers = row
        else:
            data.append(row)
    print(headers)
    print(data)

    for idx, row in enumerate(data):
        document = Document('./form_letter.docx')
        print(idx, row)
        for paragraph in document.paragraphs:
            if paragraph.text in headers:
                paragraph.text = row[headers.index(paragraph.text)]
                print(paragraph.text)
            elif paragraph.text == 'Dear name':
                paragraph.text = 'Dear {}'.format(row[headers.index('name')])
                print(paragraph.text)
            elif 'US$ amount' in paragraph.text:
                text = paragraph.text
                text = text.replace(
                    "amount",
                    str(row[headers.index('amount')])
                )
                paragraph.text = text
                print(paragraph.text)
            paragraph.style.font.name = 'Arial Narrow'
        document.save('./build/letter_{}.docx'.format(row[headers.index('name')]))
'''
    for idx, row in enumerate(reader):
        print(idx, row)
'''

