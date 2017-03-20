import csv
from docx import Document

# Read the file and put it into an array
with open('data_investor_new.csv', newline='') as csvfile:
    document = Document('./approval_letter.docx')
    reader = csv.reader(csvfile)
    data = []
    for idx, row in enumerate(reader):
        if idx == 0:
            headers = row
        else:
            data.append(row)
    #print(headers)
    #print(data)

    for idx, row in enumerate(data):
        document = Document('./approval_letter.docx')
        #print(idx, row)
        for paragraph in document.paragraphs:
            text = paragraph.text
            if 'title' in paragraph.text:
                for run in paragraph.runs:
                    if run.text == 'title':
                        if (row[headers.index('institution')] == 'TRUE'):
                            run.text = 'Sirs'
                        else:
                            run.text = str(row[headers.index('title')])
            if 'I/We' in paragraph.text:
                for run in paragraph.runs:
                    if run.text == "name1" or "name1" in run.text:
                        run.text = str(row[headers.index('name')])
            if 'name1' in paragraph.text:
                text = paragraph.text
                for run in paragraph.runs:
                    if run.text == "name1" or "name1" in run.text:
                        run.text = str(row[headers.index('name')])
            if 'address1' in paragraph.text:
                text = paragraph.text
                text = text.replace(
                    "address1",
                    str(row[headers.index('address1')]).strip()
                    )
                paragraph.text = text
            if 'address2' in paragraph.text:
                text = paragraph.text
                text = text.replace(
                    "address2",
                    str(row[headers.index('address2')]).strip()
                    )
                paragraph.text = text
            if 'address3' in paragraph.text:
                text = paragraph.text
                text = text.replace(
                    "address3",
                    str(row[headers.index('address3')]).strip()
                    )
                paragraph.text = text
            if 'amount' in paragraph.text:
                amount = '{0:,}'.format(int(row[headers.index('amount')]))
                for run in paragraph.runs:
                    if "amount" in run.text:
                        run.text = run.text.replace(
                            "US$amount",
                            'US${}'.format(amount) 
                        )
                    if 'amount2' in run.text:
                        run.text = run.text.replace(
                            "amount2",
                            str(int(int(row[headers.index('amount')])/1000))
                        )
            if 'name2' in paragraph.text:
                text = paragraph.text
                for run in paragraph.runs:
                    if "name2" in run.text:
                        if (row[headers.index('institution')] == 'TRUE'):
                            run.text = run.text.replace(
                                'name2',
                                'Name:'
                                )
                            document.add_paragraph("Title:")
                            document.add_paragraph("On behalf of: {}".format(row[headers.index('name')]))
                        else:
                            run.text = run.text.replace(
                                'name2',
                                row[headers.index('name')]
                                )
            # All checks complete
        document.save('./build/letter_{}.docx'.format(row[headers.index('name')]))

