import csv
from docx import Document

# Read the file and put it into an array
with open('data_investor.csv', newline='') as csvfile:
    document = Document('./approval_letter.docx')
    reader = csv.reader(csvfile)
    data = []
    for idx, row in enumerate(reader):
        if idx == 0:
            headers = row
        else:
            data.append(row)
    #print(headers)
    #print(data)

    for idx, row in enumerate(data):
        document = Document('./approval_letter.docx')
        print(idx, row)
        for paragraph in document.paragraphs:
            text = paragraph.text
            if 'name1' in paragraph.text:
                for run in paragraph.runs:
                    if run.text == "name1":
                        run.text = str(row[headers.index('name')])
            if 'address1' in paragraph.text:
                print
                text = text.replace(
                    "address1",
                    str(row[headers.index('address1')]).strip()
                    )
                paragraph.text = text
            if 'address2' in paragraph.text:
                text = text.replace(
                    "address2",
                    str(row[headers.index('address2')]).strip()
                    )
                paragraph.text = text
            if 'address3' in paragraph.text:
                text = text.replace(
                    "address3",
                    str(row[headers.index('address3')]).strip()
                    )
                paragraph.text = text
            if 'amount' in paragraph.text:
                amount = '{0:,}'.format(int(row[headers.index('amount')]))
                text = text.replace(
                    "US$amount",
                    'US${}'.format(amount) 
                )
                paragraph.text = text
                text = text.replace(
                    "amount2",
                    str(int(int(row[headers.index('amount')])/1000))
                )
                paragraph.text = text
            if 'name2' in paragraph.text:
                for run in paragraph.runs:
                    print(run.text)
                    if "name2" in run.text:
                        if (row[headers.index('institution')] == 'TRUE'):
                            run.text = run.text.replace(
                                'name2',
                                'Name:'
                                )
                            document.add_paragraph("Title:")
                            document.add_paragraph("On behalf of: {}".format(row[headers.index('name')]))
                        else:
                            run.text = run.text.replace(
                                'name2',
                                row[headers.index('name')]
                                )
            # All checks complete
            paragraph.style.font.name = 'Arial Narrow'
        document.save('./build/letter_{}.docx'.format(row[headers.index('name')]))

